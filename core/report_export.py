from docx import Document
from docx.shared import Inches
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import tempfile
import os
from agents.citations import get_citations
from core.apa_tables import format_group_table, format_test_table


def save_fig(fig, path):
    fig.savefig(path, bbox_inches="tight")


def generate_word(rt, fig1, fig2, results, schema, audit_log, test_plan):
    temp_dir = tempfile.mkdtemp()
    doc = Document()

    doc.add_heading("Statistical Analysis Report", 1)

    doc.add_paragraph(rt["results_text"])
    doc.add_paragraph(rt["interpretation"])
    doc.add_paragraph(rt["limitations"])

    doc.add_heading("References", 2)
    citation_data = get_citations(test_plan["selected_test"])
    for c in citation_data["citations"]:
        doc.add_paragraph(c)

    doc.add_heading("Reproducibility Appendix", 2)
    for item in audit_log:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(str(schema.to_dict()))

    fig1_path = os.path.join(temp_dir, "fig1.png")
    fig2_path = os.path.join(temp_dir, "fig2.png")

    save_fig(fig1, fig1_path)
    save_fig(fig2, fig2_path)

    doc.add_picture(fig1_path, width=Inches(4))
    doc.add_picture(fig2_path, width=Inches(4))

    path = os.path.join(temp_dir, "analysis.docx")
    doc.save(path)
    return path


def generate_pdf(rt, fig1, fig2, results, schema, audit_log, test_plan):
    temp_dir = tempfile.mkdtemp()
    pdf_path = os.path.join(temp_dir, "analysis.pdf")

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(pdf_path)
    elements = []

    elements.append(Paragraph(rt["results_text"], styles["Normal"]))
    elements.append(Paragraph(rt["interpretation"], styles["Normal"]))

    citation_data = get_citations(test_plan["selected_test"])
    for c in citation_data["citations"]:
        elements.append(Paragraph(c, styles["Italic"]))

    fig1_path = os.path.join(temp_dir, "fig1.png")
    fig2_path = os.path.join(temp_dir, "fig2.png")

    save_fig(fig1, fig1_path)
    save_fig(fig2, fig2_path)

    elements.append(Image(fig1_path, 300, 200))
    elements.append(Image(fig2_path, 300, 200))

    doc.build(elements)
    return pdf_path