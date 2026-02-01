from docx import Document
from docx.shared import Inches
import tempfile
import os


def generate_word_report(
    test_plan,
    results,
    report_text,
    figures
):
    doc = Document()

    # -----------------------
    # Title
    # -----------------------

    doc.add_heading("Statistical Analysis Report", level=1)

    # -----------------------
    # Method
    # -----------------------

    doc.add_heading("Statistical Method", level=2)
    doc.add_paragraph(f"Test used: {test_plan['selected_test']}")
    doc.add_paragraph("Assumptions:")
    for a in test_plan["assumptions"]:
        doc.add_paragraph(a, style="List Bullet")

    # -----------------------
    # Results text
    # -----------------------

    doc.add_heading("Results", level=2)
    doc.add_paragraph(report_text["results_text"])

    doc.add_heading("Interpretation", level=3)
    doc.add_paragraph(report_text["interpretation"])

    doc.add_heading("Limitations", level=3)
    doc.add_paragraph(report_text["limitations"])

    # -----------------------
    # Numeric table
    # -----------------------

    doc.add_heading("Statistical Output", level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List"

    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"

    for k, v in results.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # -----------------------
    # Figures
    # -----------------------

    doc.add_heading("Figures", level=2)

    for fig_path in figures:
        doc.add_picture(fig_path, width=Inches(4))
        doc.add_paragraph("")

    # -----------------------
    # Save temp file
    # -----------------------

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "statistical_report.docx")
    doc.save(file_path)

    return file_path